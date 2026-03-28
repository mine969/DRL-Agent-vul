from copy import deepcopy
from pathlib import Path
import zipfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parent
STRICT_TEMPLATE_PATH = BASE_DIR / "final draft.docx"
TRANSITIONAL_TEMPLATE_PATH = BASE_DIR / "final draft_transitional.docx"
SOURCE_DRAFT_PATH = BASE_DIR / "Research_Paper_Draft.docx"
ARCHITECTURE_FIGURE_PATH = BASE_DIR / "D3QN_vuln_finder_reviewer_fixed.png"
TRAINING_FIGURE_PATH = BASE_DIR / "training_curve.png"
OUTPUT_PATH = BASE_DIR / "final draft_completed_submit_ready.docx"
FALLBACK_OUTPUT_PATH = BASE_DIR / "final draft_completed_submit_ready_v2.docx"

TITLE_TEXT = "Deep Reinforcement Learning Vulnerability Scanner for Modern Web Applications"
KEYWORDS_TEXT = (
    "Keywords- deep reinforcement learning, web vulnerability scanning, reinforcement learning, "
    "autonomous penetration testing, web application security"
)

AUTHOR_LINES = [
    ("Hein Htet Zaw", False),
    ("Digital Innovative Technology", True),
    ("Rangsit University", True),
    ("Bangkok, Pathum Thani", False),
    ("devildog.kk@gmail.com", False),
]

REFERENCE_ENTRIES = [
    "[1] R. Singh, M. K. Gupta, D. R. Patil, and S. M. Patil, \"Analysis of Web Application Vulnerabilities using Dynamic Application Security Testing,\" in Proc. IEEE 9th Int. Conf. Convergence in Technology (I2CT), 2024, pp. 1-6, doi: 10.1109/I2CT61223.2024.10543484.",
    "[2] R. Sri Devi and M. Mohan Kumar, \"Testing for Security Weakness of Web Applications using Ethical Hacking,\" in Proc. 4th Int. Conf. Trends in Electronics and Informatics (ICOEI), 2020, pp. 354-361, doi: 10.1109/ICOEI48184.2020.9143018.",
    "[3] C. Mainka, J. Somorovsky, and J. Schwenk, \"Penetration Testing Tool for Web Services Security,\" in 2012 IEEE Eighth World Congress on Services, 2012, pp. 163-170, doi: 10.1109/SERVICES.2012.7.",
    "[4] V. Sujatha, K. Lakshmi Prasanna, K. Niharika, V. Charishma, and K. Bhavya Sai, \"Network Intrusion Detection using Deep Reinforcement Learning,\" in Proc. 7th Int. Conf. Computing Methodologies and Communication (ICCMC), 2023, pp. 1146-1150, doi: 10.1109/ICCMC56507.2023.10083673.",
    "[5] V. Mnih et al., \"Human-level control through deep reinforcement learning,\" Nature, vol. 518, no. 7540, pp. 529-533, 2015, doi: 10.1038/nature14236.",
    "[6] T. Schaul, J. Quan, I. Antonoglou, and D. Silver, \"Prioritized Experience Replay,\" arXiv:1511.05952, 2015. [Online]. Available: https://arxiv.org/abs/1511.05952.",
    "[7] M. C. Ghanem and T. M. Chen, \"Reinforcement learning for efficient network penetration testing,\" Information, vol. 11, no. 1, Art. no. 6, 2020, doi: 10.3390/info11010006.",
    "[8] Anonymous, \"Pentest-R1: Towards Autonomous Penetration Testing Reasoning Optimized via Two-Stage Reinforcement Learning,\" arXiv, 2024.",
    "[9] H. Al Shaikh, S. Saha, K. Zamiri Azar, F. Farahmandi, M. Tehranipoor, and F. Rahman, \"Re-Pen: Reinforcement Learning-Enforced Penetration Testing for SoC Security Verification,\" IEEE Trans. Very Large Scale Integr. (VLSI) Syst., vol. 33, no. 3, pp. 853-866, 2025, doi: 10.1109/TVLSI.2024.3510682.",
    "[10] S. Zhou, J. Liu, Y. Lu, J. Yang, Y. Zhang, B. Lin, X. Zhong, and S. Hu, \"SCRIPT: A Scalable Continual Reinforcement Learning Framework for Autonomous Penetration Testing,\" Expert Syst. Appl., vol. 285, Art. no. 127827, 2025, doi: 10.1016/j.eswa.2025.127827.",
    "[11] J. Liu, Y. Zhang, S. Zhou, J. Yang, Y. Lu, and X. Zhong, \"Autonomous penetration testing using reinforcement learning: A review and perspectives,\" Expert Syst. Appl., vol. 300, Art. no. 130219, 2026, doi: 10.1016/j.eswa.2025.130219.",
    "[12] N. Singh, V. Meherhomji, and B. R. Chandavarkar, \"Automated versus Manual Approach of Web Application Penetration Testing,\" in Proc. 11th Int. Conf. Computing, Communication and Networking Technologies (ICCCNT), 2020, pp. 1-6, doi: 10.1109/ICCCNT49239.2020.9225385.",
    "[13] D.-Y. Kao, Y.-Y. Chen, and F.-C. Tsai, \"Hacking Tool Identification in Penetration Testing,\" in Proc. 22nd Int. Conf. Advanced Communication Technology (ICACT), 2020, pp. 256-261, doi: 10.23919/ICACT48636.2020.9061401.",
    "[14] A. Chowdhary, D. Huang, J. S. Mahendran, D. Romo, Y. Deng, and A. Sabur, \"Autonomous Security Analysis and Penetration Testing,\" in Proc. 16th Int. Conf. Mobility, Sensing and Networking (MSN), 2020, pp. 508-515, doi: 10.1109/MSN50589.2020.00086.",
]

STRICT_TO_TRANSITIONAL = {
    "http://purl.oclc.org/ooxml/officeDocument/relationships/officeDocument": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/styles": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/settings": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/webSettings": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/fontTable": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/theme": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/numbering": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/footer": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/image": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/hyperlink": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/footnotes": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/endnotes": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/comments": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/extendedProperties": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties",
    "http://purl.oclc.org/ooxml/officeDocument/relationships": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "http://purl.oclc.org/ooxml/officeDocument/math": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "http://purl.oclc.org/ooxml/wordprocessingml/main": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def convert_template_to_transitional():
    with zipfile.ZipFile(STRICT_TEMPLATE_PATH, "r") as zin, zipfile.ZipFile(
        TRANSITIONAL_TEMPLATE_PATH, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8")
                for source, target in STRICT_TO_TRANSITIONAL.items():
                    text = text.replace(source, target)
                data = text.encode("utf-8")
            zout.writestr(item, data)


def ensure_style(document, style_name, base_name=None):
    styles = document.styles
    if style_name in styles:
        return styles[style_name]
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    if base_name and base_name in styles:
        style.base_style = styles[base_name]
    return style


def set_style_font(style, name=None, size=None, bold=None, italic=None):
    if name:
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:ascii"), name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def ensure_missing_styles(document):
    compact = ensure_style(document, "Compact", base_name="Body Text")
    set_style_font(compact, name="Times New Roman", size=10)
    compact.paragraph_format.first_line_indent = Pt(0)
    compact.paragraph_format.left_indent = Pt(0)
    compact.paragraph_format.space_before = Pt(0)
    compact.paragraph_format.space_after = Pt(0)
    compact.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    captioned_figure = ensure_style(document, "Captioned Figure", base_name="Body Text")
    set_style_font(captioned_figure, name="Times New Roman", size=10)
    captioned_figure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    captioned_figure.paragraph_format.first_line_indent = Pt(0)
    captioned_figure.paragraph_format.space_before = Pt(3)
    captioned_figure.paragraph_format.space_after = Pt(0)

    image_caption = ensure_style(document, "Image Caption", base_name="figure caption")
    set_style_font(image_caption, name="Times New Roman", size=8)
    image_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_caption.paragraph_format.space_before = Pt(3)
    image_caption.paragraph_format.space_after = Pt(6)

    abstract_title = ensure_style(document, "Abstract Title", base_name="Normal")
    set_style_font(abstract_title, name="Times New Roman", size=10, bold=True)
    abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_title.paragraph_format.space_before = Pt(0)
    abstract_title.paragraph_format.space_after = Pt(0)

    source_code = ensure_style(document, "Source Code", base_name="Normal")
    set_style_font(source_code, name="Courier New", size=7)
    source_code.paragraph_format.space_before = Pt(0)
    source_code.paragraph_format.space_after = Pt(0)
    source_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_continuous_section_break_after(paragraph, columns_for_previous_section=1):
    body_sectpr = paragraph._parent._element.sectPr
    new_sectpr = deepcopy(body_sectpr)

    cols = new_sectpr.xpath("./w:cols")
    if cols:
        cols_element = cols[0]
    else:
        cols_element = OxmlElement("w:cols")
        new_sectpr.append(cols_element)
    cols_element.set(qn("w:num"), str(columns_for_previous_section))
    cols_element.set(qn("w:space"), str(int(Inches(0.25).twips)))

    type_element = new_sectpr.find(qn("w:type"))
    if type_element is None:
        type_element = OxmlElement("w:type")
        new_sectpr.insert(0, type_element)
    type_element.set(qn("w:val"), "continuous")

    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(new_sectpr)


def set_section_columns(section, count, space_inches):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_element = cols[0]
    else:
        cols_element = OxmlElement("w:cols")
        sect_pr.append(cols_element)
    cols_element.set(qn("w:num"), str(count))
    cols_element.set(qn("w:space"), str(int(Inches(space_inches).twips)))


def set_run_font(run, name="Times New Roman", size=9, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def configure_front_matter(document, source_draft):
    body_children = list(document._body._element)
    kept_children = [body_children[0], body_children[4], body_children[10], body_children[11], document._body._element.sectPr]
    for child in list(document._body._element):
        if child not in kept_children and child.tag != qn("w:sectPr"):
            document._body._element.remove(child)

    title_paragraph, author_paragraph, abstract_paragraph, keywords_paragraph = document.paragraphs[:4]

    clear_paragraph(title_paragraph)
    title_run = title_paragraph.add_run(TITLE_TEXT)
    set_run_font(title_run, size=24)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    clear_paragraph(author_paragraph)
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (line_text, italic) in enumerate(AUTHOR_LINES):
        if index:
            author_paragraph.add_run().add_break()
        run = author_paragraph.add_run(line_text)
        set_run_font(run, size=9, italic=italic)

    clear_paragraph(abstract_paragraph)
    abstract_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_run = abstract_paragraph.add_run("Abstract-")
    body_run = abstract_paragraph.add_run(f" {source_draft.paragraphs[2].text.strip()}")
    set_run_font(label_run, size=9, bold=True)
    set_run_font(body_run, size=9)

    clear_paragraph(keywords_paragraph)
    keywords_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_run = keywords_paragraph.add_run(KEYWORDS_TEXT)
    set_run_font(kw_run, size=9, italic=True)


def append_body_element(destination_doc, element):
    destination_doc._body._element.insert(len(destination_doc._body._element) - 1, deepcopy(element))


def add_picture_block(document, image_path, caption_text, width_inches=3.15):
    paragraph = document.add_paragraph(style="Captioned Figure")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption = document.add_paragraph(caption_text, style="figure caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_caption(document, caption_text):
    paragraph = document.add_paragraph(caption_text, style="table head")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_references(document):
    heading = document.add_paragraph("REFERENCES", style="Heading 5")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for entry in REFERENCE_ENTRIES:
        paragraph = document.add_paragraph(entry, style="references")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def copy_main_content(destination_doc, source_doc):
    body_items = list(source_doc._body._element.iterchildren())
    start_index = None
    end_index = None

    for index, child in enumerate(body_items):
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, source_doc._body)
            text = " ".join(paragraph.text.split())
            if text == "I. INTRODUCTION":
                start_index = index
            if text == "REFERENCES":
                end_index = index
                break

    if start_index is None or end_index is None:
        raise ValueError("Could not locate draft body range")

    results_table_added = False
    figure_count = 0

    for child in body_items[start_index:end_index]:
        if isinstance(child, CT_Tbl):
            if not results_table_added:
                results_table_added = True
            else:
                add_table_caption(destination_doc, "TABLE I. AVERAGE CONFIRMED FINDINGS ACROSS FIVE RUNS")
            append_body_element(destination_doc, child)
            continue

        paragraph = Paragraph(child, source_doc._body)
        style_name = paragraph.style.name

        if style_name == "Captioned Figure":
            figure_count += 1
            if figure_count == 1:
                add_picture_block(
                    destination_doc,
                    ARCHITECTURE_FIGURE_PATH,
                    "Fig. 1. System architecture of the proposed RL-based web vulnerability scanner.",
                )
            elif figure_count == 2:
                add_picture_block(
                    destination_doc,
                    TRAINING_FIGURE_PATH,
                    "Fig. 2. Extended D3QN agent training progression across 10,000 episodes.",
                )
            continue

        if style_name == "Image Caption":
            continue

        append_body_element(destination_doc, child)


def generate_completed_final_draft():
    convert_template_to_transitional()

    template_doc = Document(str(TRANSITIONAL_TEMPLATE_PATH))
    source_doc = Document(str(SOURCE_DRAFT_PATH))

    ensure_missing_styles(template_doc)
    configure_front_matter(template_doc, source_doc)
    copy_main_content(template_doc, source_doc)
    add_references(template_doc)
    add_continuous_section_break_after(template_doc.paragraphs[3], columns_for_previous_section=1)
    set_section_columns(template_doc.sections[-1], 2, 0.25)

    try:
        template_doc.save(str(OUTPUT_PATH))
        return OUTPUT_PATH
    except PermissionError:
        template_doc.save(str(FALLBACK_OUTPUT_PATH))
        return FALLBACK_OUTPUT_PATH


def main():
    output = generate_completed_final_draft()
    print(output)


if __name__ == "__main__":
    main()
