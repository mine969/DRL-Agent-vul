import copy
import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parent
SOURCE_PATH = BASE_DIR / "Research_Paper_Draft.docx"
OUTPUT_PATH = BASE_DIR / "Research_Paper_IEEE_Formatted_v3.docx"
AUTHOR_DETAILS_PATH = BASE_DIR / "paper_author_details.json"

DEFAULT_AUTHOR_BLOCKS = [
    {
        "name": f"Author {index} Name Surname",
        "department": "Department / School",
        "organization": "University / Organization",
        "location": "City, Country",
        "email": f"author{index}@example.com",
    }
    for index in range(1, 7)
]

KEYWORDS_TEXT = (
    "Keywords- deep reinforcement learning, web vulnerability scanning, "
    "reinforcement learning, autonomous penetration testing, web application security"
)


def int_to_roman(value):
    numerals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    result = []
    for integer, numeral in numerals:
        while value >= integer:
            result.append(numeral)
            value -= integer
    return "".join(result)


def load_author_blocks(author_details_path=None):
    path = Path(author_details_path) if author_details_path else AUTHOR_DETAILS_PATH
    if path.exists():
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            data = data.get("authors", [])
        blocks = []
        for item in data:
            if not isinstance(item, dict):
                continue
            name = (item.get("name") or "").strip()
            if not name:
                continue
            blocks.append(
                {
                    "name": name,
                    "department": (item.get("department") or "").strip(),
                    "organization": (item.get("organization") or "").strip(),
                    "location": (item.get("location") or "").strip(),
                    "email": (item.get("email") or item.get("orcid") or "").strip(),
                }
            )
        if blocks:
            return blocks
    return copy.deepcopy(DEFAULT_AUTHOR_BLOCKS)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph, text):
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def set_run_font(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def apply_run_format_to_paragraph(paragraph, name="Times New Roman", size=10, bold=False, italic=False):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=bold, italic=italic)


def ensure_style(document, style_name, base_name=None):
    styles = document.styles
    if style_name in styles:
        return styles[style_name]
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    if base_name and base_name in styles:
        style.base_style = styles[base_name]
    return style


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.8))
    paragraph._p.addnext(table._tbl)
    return table


def configure_document_defaults(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)

    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)


def clean_document_metadata(document, title_text):
    props = document.core_properties
    props.title = title_text
    props.author = ""
    props.last_modified_by = ""
    props.category = "Conference Paper"
    props.comments = ""
    props.content_status = "Final Draft"
    props.identifier = ""
    props.keywords = KEYWORDS_TEXT.replace("Keywords-", "", 1).strip()
    props.language = "en-US"
    props.subject = "Deep reinforcement learning for web vulnerability scanning"


def set_section_columns(section, count, space_inches):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(int(Inches(space_inches).twips)))


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = tbl_borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tbl_borders.append(edge_el)
        edge_el.set(qn("w:val"), "nil")


def add_continuous_section_break_after(paragraph, columns_for_previous_section=1):
    body_sectpr = paragraph._parent._element.sectPr
    new_sectpr = copy.deepcopy(body_sectpr)

    cols = new_sectpr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        new_sectpr.append(cols_el)
    cols_el.set(qn("w:num"), str(columns_for_previous_section))
    cols_el.set(qn("w:space"), str(int(Inches(0.25).twips)))

    type_el = new_sectpr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectpr.insert(0, type_el)
    type_el.set(qn("w:val"), "continuous")

    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(new_sectpr)


def configure_styles(document):
    body_like = ["Body Text", "First Paragraph", "Compact", "Normal"]
    for style_name in body_like:
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Body Text" in document.styles:
        document.styles["Body Text"].paragraph_format.first_line_indent = Inches(0.17)
    if "First Paragraph" in document.styles:
        document.styles["First Paragraph"].paragraph_format.first_line_indent = Inches(0)
    if "Compact" in document.styles:
        document.styles["Compact"].paragraph_format.first_line_indent = Inches(0)

    if "Heading 1" in document.styles:
        style = document.styles["Heading 1"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)

    if "Heading 2" in document.styles:
        style = document.styles["Heading 2"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)

    if "Heading 3" in document.styles:
        style = document.styles["Heading 3"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.font.italic = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)

    if "Image Caption" in document.styles:
        style = document.styles["Image Caption"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(8)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)

    if "Source Code" in document.styles:
        style = document.styles["Source Code"]
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        style.font.size = Pt(7)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    table_caption = ensure_style(document, "IEEE Table Caption", base_name="Normal")
    table_caption.font.name = "Times New Roman"
    table_caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    table_caption.font.size = Pt(8)
    table_caption.font.bold = True
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.space_before = Pt(6)
    table_caption.paragraph_format.space_after = Pt(3)

    references_style = ensure_style(document, "IEEE References", base_name="Normal")
    references_style.font.name = "Times New Roman"
    references_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    references_style.font.size = Pt(8)
    references_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    references_style.paragraph_format.first_line_indent = Inches(-0.25)
    references_style.paragraph_format.left_indent = Inches(0.25)
    references_style.paragraph_format.space_before = Pt(0)
    references_style.paragraph_format.space_after = Pt(0)
    references_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def populate_author_cell(cell, block):
    cell.width = Inches(2.2)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    lines = [
        (block.get("name", ""), {"size": 10, "bold": True}),
        (block.get("department", ""), {"size": 9, "italic": True}),
        (block.get("organization", ""), {"size": 9, "italic": True}),
        (block.get("location", ""), {"size": 9}),
        (block.get("email", ""), {"size": 9}),
    ]

    printed = False
    for text, options in lines:
        if not text:
            continue
        if printed:
            paragraph.add_run().add_break()
        run = paragraph.add_run(text)
        set_run_font(run, **options)
        printed = True


def insert_author_block(title_paragraph, author_paragraph, author_blocks):
    total_authors = max(1, len(author_blocks))
    if total_authors <= 3:
        rows = 1
        cols = total_authors
    else:
        rows = 2
        cols = math.ceil(total_authors / 2)

    table = insert_table_after(title_paragraph, rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    index = 0
    for row in table.rows:
        row.height = Inches(0.95)
        for cell in row.cells:
            clear_paragraph(cell.paragraphs[0])
            if index < total_authors:
                populate_author_cell(cell, author_blocks[index])
                index += 1

    remove_paragraph(author_paragraph)


def format_title_block(document, author_blocks):
    title = document.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    apply_run_format_to_paragraph(title, size=24)

    author = document.paragraphs[1]
    abstract = document.paragraphs[2]
    abstract_text = abstract.text.strip()
    if not abstract_text.lower().startswith("abstract-"):
        abstract_text = f"Abstract- {abstract_text}"
    clear_paragraph(abstract)
    label_run = abstract.add_run("Abstract-")
    text_run = abstract.add_run(f" {abstract_text.replace('Abstract-', '', 1).strip()}")
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_before = Pt(0)
    abstract.paragraph_format.space_after = Pt(4)
    set_run_font(label_run, size=9, bold=True)
    set_run_font(text_run, size=9)

    keywords = document.paragraphs[3]
    clear_paragraph(keywords)
    kw_label = keywords.add_run("Keywords-")
    kw_text = keywords.add_run(f" {KEYWORDS_TEXT.replace('Keywords-', '', 1).strip()}")
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.space_before = Pt(0)
    keywords.paragraph_format.space_after = Pt(6)
    set_run_font(kw_label, size=9, italic=True)
    set_run_font(kw_text, size=9)

    insert_author_block(title, author, author_blocks)


def insert_paragraph_before_table(table, style_name, text):
    new_p = OxmlElement("w:p")
    table._element.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)
    return paragraph


def format_tables(document):
    if not document.tables:
        return

    content_tables = list(document.tables)[1:]

    caption_index = 1
    for table in content_tables:
        header_text = " | ".join(cell.text.strip() for cell in table.rows[0].cells).lower()
        if "target" in header_text or "website" in header_text:
            caption_body = "REPEATED EVALUATION RESULTS ACROSS FIVE TARGETS"
        elif "algorithm" in header_text:
            caption_body = "PHASE-BASED EXTENDED D3QN TRAINING PROCESS"
        else:
            caption_body = "RESULTS SUMMARY"
        caption = f"TABLE {int_to_roman(caption_index)}. {caption_body}"
        caption_index += 1
        insert_paragraph_before_table(table, "IEEE Table Caption", caption)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        if paragraph.style.name == "Source Code":
                            set_run_font(run, name="Courier New", size=7)
                        else:
                            set_run_font(run, size=8)
                            if row_index == 0:
                                run.font.bold = True


def resize_figures(document):
    for shape in document.inline_shapes:
        if shape.width > Inches(3.2):
            ratio = shape.height / shape.width
            shape.width = Inches(3.2)
            shape.height = int(shape.width * ratio)

    for paragraph in document.paragraphs:
        if paragraph.style.name == "Captioned Figure":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(0)
        elif paragraph.style.name == "Image Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_run_format_to_paragraph(paragraph, size=8)


def split_references(document):
    heading_index = None
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().upper() == "REFERENCES":
            heading_index = idx
            break
    if heading_index is None or heading_index + 1 >= len(document.paragraphs):
        return

    refs_paragraph = document.paragraphs[heading_index + 1]
    refs_text = refs_paragraph.text.strip()
    refs = [item.strip() for item in re.split(r"(?=\[\d+\])", refs_text) if item.strip()]
    if not refs:
        return

    set_paragraph_text(refs_paragraph, refs[0])
    refs_paragraph.style = document.styles["IEEE References"]
    apply_run_format_to_paragraph(refs_paragraph, size=8)

    for ref in refs[1:]:
        new_para = document.add_paragraph(ref)
        new_para.style = document.styles["IEEE References"]
        apply_run_format_to_paragraph(new_para, size=8)


def remove_extra_blank_paragraphs(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            continue
        if paragraph.style.name in {"Normal", "Body Text", "First Paragraph"}:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)


def generate(source_path=None, output_path=None, author_details_path=None):
    resolved_source = Path(source_path) if source_path else SOURCE_PATH
    resolved_output = Path(output_path) if output_path else OUTPUT_PATH
    author_blocks = load_author_blocks(author_details_path)

    document = Document(str(resolved_source))
    configure_document_defaults(document)
    configure_styles(document)
    clean_document_metadata(document, document.paragraphs[0].text.strip())
    format_title_block(document, author_blocks)

    add_continuous_section_break_after(document.paragraphs[3], columns_for_previous_section=1)
    set_section_columns(document.sections[-1], 2, 0.25)

    format_tables(document)
    resize_figures(document)
    split_references(document)
    remove_extra_blank_paragraphs(document)

    document.save(str(resolved_output))
    return resolved_output


if __name__ == "__main__":
    output = generate()
    print(output)
