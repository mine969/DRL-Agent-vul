import json
from pathlib import Path

from docx import Document


RESULTS_PATH = Path("research/results/autonomous_scan_average_findings.json")


HUMANIZED_PARAGRAPHS = [
    "To make the evaluation more trustworthy, we did not rely on a single scan. Instead, each mock website was tested five times with autonomous_scan.py. In every run, the scanner auto-loaded the improved_mock_ep10000.pth checkpoint and used the same scan settings (--depth 10 --intensity 2), which gave us a fair basis for averaging the confirmed findings.",
    "Table I shows the average number of confirmed findings across those five runs, together with the matching detection rate for each vulnerability category.",
    "What stands out most is not variability, but consistency. Across all five runs, the scanner ended with zero confirmed findings after its false-positive filtering phase. In other words, the average finding count stayed at 0.0 for every listed category, and every averaged detection rate remained at 0.0% under this exact configuration.",
    "That result should not be read as proof that the mock applications were safe. We already know these targets contain intentionally planted vulnerabilities. A more reasonable interpretation is that the current autonomous setup is still too cautious: it can crawl, inspect, and generate low-confidence signals, but it is not yet reliably converting those signals into confirmed exploit detections. From a research perspective, that points to the next improvement area very clearly: the model, exploration budget, and confirmation logic all need stronger tuning before the scanner can deliver dependable autonomous findings.",
]


BACHELOR_PARAGRAPHS = [
    "To make the results more reliable, each mock website was scanned five times with autonomous_scan.py instead of only once. Every run used the same automatically loaded checkpoint, improved_mock_ep10000.pth, and the same scan settings (--depth 10 --intensity 2). This made the comparison fair across all five targets.",
    "Table I shows the average number of confirmed findings from those five runs.",
    "The repeated tests show a very clear pattern. The scanner did not return confirmed findings on any of the five mock websites after the false-positive filtering step. Because of this, the average finding count stayed at 0.0 for every category, and the detection rate also stayed at 0.0% in this experiment.",
    "This does not mean that the websites had no vulnerabilities. The vulnerabilities were still present in the mock targets. It means that the current scan setup was too weak or too limited to confirm them. In simple terms, the scanner could move around the websites, but it still needs better tuning, deeper exploration, and stronger validation before it can find vulnerabilities in a reliable way.",
]


def load_rows():
    data = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    ordered_targets = ["ecommerce", "social", "banking", "blog", "fileshare"]
    rows = []
    for target_key in ordered_targets:
        target = data["targets"].get(target_key)
        if not target:
            continue
        for row in target["rows"]:
            rows.append(
                [
                    target["name"],
                    row["category"],
                    str(row["total_existing"]),
                    str(row["average_detected"]),
                    row["detection_rate"],
                    row["severity"],
                ]
            )
    return rows


def find_heading_index(doc: Document, heading_text: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading_text:
            return index
    raise ValueError(f"Heading not found: {heading_text}")


def find_next_heading_index(doc: Document, start_index: int, heading_prefix: str) -> int:
    for index in range(start_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[index].text.strip().startswith(heading_prefix):
            return index
    return len(doc.paragraphs)


def update_paragraph_block(doc: Document, heading_text: str, paragraphs: list[str]):
    start = find_heading_index(doc, heading_text)
    end = find_next_heading_index(doc, start, "V.")
    target_paragraphs = doc.paragraphs[start + 1 : end]
    for index, paragraph in enumerate(target_paragraphs):
        paragraph.text = paragraphs[index] if index < len(paragraphs) else ""


def resize_table(table, required_rows: int):
    while len(table.rows) < required_rows:
        table.add_row()
    while len(table.rows) > required_rows:
        table._tbl.remove(table.rows[-1]._tr)


def update_results_table(doc: Document, rows: list[list[str]]):
    table = doc.tables[1]
    resize_table(table, len(rows) + 1)
    header = [
        "Website",
        "Vulnerability Type",
        "Total Existing",
        "Average Findings (5 Runs)",
        "Detection Rate",
        "Severity",
    ]
    for col_index, value in enumerate(header):
        table.rows[0].cells[col_index].text = value
    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            table.rows[row_index].cells[col_index].text = value


def update_docx(path: str, paragraphs: list[str], rows: list[list[str]]):
    doc = Document(path)
    update_paragraph_block(doc, "IV. EVALUATION AND RESULTS", paragraphs)
    update_results_table(doc, rows)
    doc.save(path)


def main():
    rows = load_rows()
    update_docx("research/Research_Paper_Draft.docx", HUMANIZED_PARAGRAPHS, rows)
    update_docx("research/Bachelor_Paper_Draft.docx", BACHELOR_PARAGRAPHS, rows)


if __name__ == "__main__":
    main()
