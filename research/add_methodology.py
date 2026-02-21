from docx import Document

doc = Document('IEEE onference-template-a4.docx')

doc.add_heading('III. METHODOLOGY', level=1)
doc.add_paragraph('This section details the design, architecture, and implementation of the proposed Reinforcement Learning (RL) based autonomous vulnerability scanner. The methodology is structured around the core components of the system: the system architecture, the formulation of vulnerability scanning as a Markov Decision Process (MDP), the design of the Deep Q-Network (DQN) agent, and the implementation of the simulation environment.')

doc.add_heading('A. System Architecture', level=2)
doc.add_paragraph('The autonomous vulnerability scanner is built upon a modular architecture designed to decouple the AI decision-making process from the underlying HTTP execution engine. The system comprises three primary modules:')
doc.add_paragraph('1) The Target Environment (WebSecurityGym): A custom-built OpenAI Gymnasium-compatible environment that simulates realistic web applications. It translates HTTP responses into numerical state vectors and processes the agent\'s actions into actual security payloads.')
doc.add_paragraph('2) The RL Agent (DQNAgent): A Deep Q-Network implementation utilizing a Dueling network architecture with experience replay. This module acts as the "brain," learning to sequence attacks to maximize the discovery of vulnerabilities.')
doc.add_paragraph('3) The Orchestrator (SecurityAuditor & WebsiteExplorer): A functional wrapper that acts as the "body." It handles initial reconnaissance (crawling the target URL, extracting forms, and finding endpoints) and executes the actions selected by the agent against the target infrastructure.')
doc.add_paragraph('4) Mock Applications: A suite of deliberately vulnerable web applications (e-commerce, social media, banking, blog, file sharing) used exclusively for safe, controlled agent training and baseline evaluation.')

doc.add_heading('B. Formulation of the RL Environment (MDP)', level=2)
doc.add_paragraph('To apply reinforcement learning to security testing, the interaction between the scanner and the web application is modeled as a Markov Decision Process (MDP), defined by the tuple (S, A, P, R, \gamma).')
doc.add_heading('1) State Representation (S)', level=3)
doc.add_paragraph('The state vector represents the agent\'s current understanding of the target web application and the outcome of the most recent interaction. The feature vector is heavily engineered to provide context without combinatorial explosion. A 15-dimensional state representation is utilized, including:')
doc.add_paragraph('- HTTP Response Metrics: Status codes (e.g., 200, 403, 500) normalized to standard categories, response time (to detect time-based vulnerabilities like SQL injection), and response size variability.')
doc.add_paragraph('- Structural Indicators: The presence of new forms, input fields, or reflection of user input in the DOM.')
doc.add_paragraph('- Security Context: Flags indicating the detection of Web Application Firewalls (WAFs), rate-limiting mechanisms, or error disclosures (e.g., SQL syntax errors).')

doc.add_heading('2) Action Space (A)', level=3)
doc.add_paragraph('The action space defines the set of security testing techniques available to the agent. The system utilizes a discrete action space comprising up to 150 distinct operations, categorized into three phases aligned with the Cyber Kill Chain:')
doc.add_paragraph('- Reconnaissance Actions (0-29): Probing for common endpoints (e.g., /admin, /api), enumerating parameters, and identifying technologies.')
doc.add_paragraph('- Vulnerability Assessment Actions (30-69): Sending generic syntax-breaking payloads (e.g., \' OR 1=1--, <script>) to elicit anomalous behaviors.')
doc.add_paragraph('- Exploitation Actions (70-149): Deploying specific, complex payloads targeting Cross-Site Scripting (XSS), SQL Injection (SQLi), Local File Inclusion (LFI), and Insecure Direct Object References (IDOR).')

doc.add_heading('C. Agent Design and Training Strategy', level=2)
doc.add_paragraph('The intelligence of the scanner is powered by a Double Dueling Deep Q-Network (D3QN). This architecture addresses the overestimation bias inherent in standard DQN while efficiently learning the value of states independent of the actions taken.')
doc.add_heading('1) Neural Network Architecture', level=3)
doc.add_paragraph('The agent utilizes a multi-layer perceptron (MLP) architecture. The input layer matches the state dimension, followed by a shared feature extraction layer (typically 256 and 128 units). The dueling architecture then splits into two streams:')
doc.add_paragraph('- A Value Stream V(s) estimating the inherent value of being in a particular state.')
doc.add_paragraph('- An Advantage Stream A(s, a) estimating the relative advantage of taking each specific action in that state.')

doc.add_heading('2) Reward Design (R)', level=3)
doc.add_paragraph('The reward function is meticulously shaped to encourage vulnerability discovery while penalizing noisy or repetitive behavior. It incorporates:')
doc.add_paragraph('- Positive Rewards: High rewards for verifying a vulnerability (e.g., +1.0) and finding Capture The Flag (CTF) markers (+5.0). Minor rewards are granted for discovering new endpoints.')
doc.add_paragraph('- Negative Penalities: A small negative step penalty (e.g., -0.01) is applied to encourage efficiency. Severe penalties are applied for triggering WAF rules or rate limits (-0.1), directly discouraging brute-force behavior.')

doc.add_heading('3) Phase-Based Learning', level=3)
doc.add_paragraph('To optimize training instability, the environment implements a Phase-Based Learning algorithm. The agent is initially restricted to reconnaissance actions. As it successfully discovers targets, it dynamically unlocks assessment and exploitation actions, providing a guided learning curriculum over the course of thousands of episodes.')

doc.add_heading('D. Evaluation Metrics', level=2)
doc.add_paragraph('To measure the efficacy of the autonomous scanner, evaluating against standard penetration testing benchmarks is crucial. Metrics include:')
doc.add_paragraph('- Vulnerability Discovery Rate: The percentage of known vulnerabilities successfully detected across the diverse mock targets.')
doc.add_paragraph('- Action Efficiency: The average number of requests required to discover a unique vulnerability, compared against traditional fuzzing or broad-spectrum scanners.')
doc.add_paragraph('- False Positive Rate: Ensuring that the agent correctly parses its state vector and does not confidently report non-exploitable anomalies.')

doc.save('IEEE_conference_template_with_methodology.docx')
print("Successfully generated IEEE_conference_template_with_methodology.docx")
